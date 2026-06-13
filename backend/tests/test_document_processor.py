"""Unit tests for app.core.document_processor.DocumentProcessor."""

from pathlib import Path

import fitz
import pytest
from docx import Document

from app.core.document_processor import DocumentProcessor


@pytest.fixture
def processor() -> DocumentProcessor:
    return DocumentProcessor()


def test_process_txt(tmp_path: Path, processor: DocumentProcessor) -> None:
    file_path = tmp_path / "notes.txt"
    file_path.write_text("Hello world.\nThis is a plain text test file.")

    sections = processor.process(str(file_path), "notes.txt")

    assert len(sections) == 1
    assert "Hello world." in sections[0]["text"]
    assert sections[0]["metadata"] == {"filename": "notes.txt", "page_number": 1, "chunk_index": 0}


def test_process_docx(tmp_path: Path, processor: DocumentProcessor) -> None:
    file_path = tmp_path / "report.docx"
    document = Document()
    document.add_paragraph("First paragraph.")
    document.add_paragraph("Second paragraph.")
    document.save(str(file_path))

    sections = processor.process(str(file_path), "report.docx")

    assert len(sections) == 1
    assert "First paragraph." in sections[0]["text"]
    assert "Second paragraph." in sections[0]["text"]
    assert sections[0]["metadata"]["filename"] == "report.docx"
    assert sections[0]["metadata"]["page_number"] == 1


def test_process_pdf(tmp_path: Path, processor: DocumentProcessor) -> None:
    file_path = tmp_path / "doc.pdf"
    pdf = fitz.open()
    page1 = pdf.new_page()
    page1.insert_text((72, 72), "Page one content")
    page2 = pdf.new_page()
    page2.insert_text((72, 72), "Page two content")
    pdf.save(str(file_path))
    pdf.close()

    sections = processor.process(str(file_path), "doc.pdf")

    assert len(sections) == 2
    assert sections[0]["metadata"]["page_number"] == 1
    assert sections[1]["metadata"]["page_number"] == 2
    assert "Page one content" in sections[0]["text"]
    assert "Page two content" in sections[1]["text"]


def test_process_unsupported_extension_raises_value_error(tmp_path: Path, processor: DocumentProcessor) -> None:
    file_path = tmp_path / "image.png"
    file_path.write_bytes(b"not a real image")

    with pytest.raises(ValueError):
        processor.process(str(file_path), "image.png")
